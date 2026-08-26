"""
File Service — Sandboxed file operations.
Security: Only allows operations within allowed directories.
"""
import csv
import os
import re
import shutil
import glob
import time
from pathlib import Path
from typing import Dict, List, Optional, Any


# Allowed root directories (user can customize)
DEFAULT_ALLOWED_ROOTS = [
    os.path.expanduser("~/Desktop"),
    os.path.expanduser("~/Documents"),
    os.path.expanduser("~/Downloads"),
]

# Dangerous paths that are NEVER allowed
BLOCKED_PATHS = [
    "C:\\Windows", "C:\\Program Files", "C:\\Program Files (x86)",
    "/usr", "/bin", "/sbin", "/etc", "/boot", "/lib",
    os.path.expanduser("~/.ssh"),
    os.path.expanduser("~/AppData/Local"),
]

# Subfolders of the data dir the AI may not touch even though the rest of
# the data dir is in its sandbox. groups/ holds the group manifests
# (tubecli.core.group_context): which agent may use which file, folder or
# sheet, and the owner's playbook notes. An agent that could write there
# through its own file_action would widen its own permissions, or plant a
# note the prompt presents as the owner's instructions — so in enforced
# mode the subtree does not exist. The human UI (enforce_roots=False) is
# not affected: the owner may inspect their own manifests.
#
# extensions_data/browser/browser_profiles is the same story one floor down.
# A profile folder is the browser's user-data-dir: its config.json holds the
# saved account (email, password, recovery address, 2FA codes), its Cookies
# database holds the live sessions, and preview_cdp.json publishes the CDP
# port of the live view — the address of a browser already logged into those
# accounts. An agent that can read that folder does not need to run a script
# to take an account, and one that can write preview_cdp.json can point its
# own group's attach at ANOTHER group's browser. Nothing AI-facing writes
# there: profiles are managed by the browser extension, and script downloads
# go to their own output_dir.
AI_PROTECTED_DATA_SUBDIRS = ["groups",
                             os.path.join("extensions_data", "browser", "browser_profiles")]

MAX_FILE_SIZE_MB = 50  # Max file size for read operations


def _safe_str(value: str) -> str:
    """Make a filesystem string safe to hand to the JSON encoder.

    On POSIX a filename may hold bytes that are not valid UTF-8, and os.listdir /
    glob hand them back as lone surrogates — PEP 383's surrogateescape. Starlette
    renders a response with json.dumps(..., ensure_ascii=False).encode("utf-8")
    (verified against the installed starlette 0.50.0) and that encoder raises
    UnicodeEncodeError on a lone surrogate, so a single oddly-named file used to
    take down the entire /list, /search and /info response it appeared in — the
    user saw a 500 for a folder that is perfectly readable, with nothing naming
    the file at fault. Windows cannot hit this (NTFS names are valid UTF-16).

    Repaired strings are for display only: the repair is lossy, so a client
    cannot use the returned path to reach that particular file. That is a real
    limitation and it is unavoidable — the name simply has no JSON spelling.

    Duplicated from cleanup._safe_str rather than imported, because this module
    is the sandbox every other part of the extension depends on and it must not
    gain an import edge to a 1,643-line module that can fail to load.
    """
    try:
        value.encode("utf-8")
        return value
    except UnicodeEncodeError:
        return value.encode("utf-8", "replace").decode("utf-8", "replace")


class FileService:
    """Sandboxed file operations service.

    Two enforcement modes share one implementation:
      enforce_roots=True  — the allowlist applies. This is what the AI (brain
                            file_action, workflow nodes, Telegram) gets.
      enforce_roots=False — the allowlist is skipped for a logged-in human using
                            the File Manager UI; they already own the machine.
                            BLOCKED_PATHS still applies in both modes: it exists
                            to stop accidents (deleting C:\\Windows, reading
                            ~/.ssh), not just untrusted callers.
    """

    def __init__(self, extra_roots: Optional[List[str]] = None, enforce_roots: bool = True):
        self.enforce_roots = enforce_roots
        self.allowed_roots = list(DEFAULT_ALLOWED_ROOTS)
        # Add data dir
        data_dir = os.environ.get("TUBECLI_DATA_DIR", "data")
        self.allowed_roots.append(os.path.abspath(data_dir))
        if extra_roots:
            self.allowed_roots.extend(extra_roots)
        # Normalize all paths
        self.allowed_roots = [os.path.normpath(r) for r in self.allowed_roots]
        # Both data roots get the protected subfolders: the one this sandbox
        # exposes (cwd-relative, like the allowlist entry above) and the one
        # the package writes the manifests under. They coincide when the
        # service runs from the repo root; a repo kept under ~/Documents and
        # started elsewhere would otherwise leave the real groups/ reachable
        # through the Documents root.
        data_roots = [os.path.abspath(data_dir)]
        try:
            from tubecli.config import DATA_DIR as _package_data_dir
            data_roots.append(os.path.abspath(str(_package_data_dir)))
        except Exception:
            pass
        self.ai_blocked = [os.path.normpath(os.path.join(root, sub))
                           for root in dict.fromkeys(data_roots)
                           for sub in AI_PROTECTED_DATA_SUBDIRS]

    @staticmethod
    def _under(path: str, root: str) -> bool:
        """Is `path` the root itself, or something inside it?

        Two traps this closes. A bare startswith() treats
        "…/Downloads_evil" as inside "…/Downloads", so the separator has to be
        part of the comparison. And Windows paths are case-insensitive, so a
        plain string compare let "c:\\windows\\system32" slip past a blocklist
        entry written as "C:\\Windows" — and rejected a legitimate
        "c:\\tubecreate-vue\\tubecli\\data" that differed only in drive case.
        os.path.normcase handles both platforms correctly.
        """
        p = os.path.normcase(os.path.normpath(path))
        r = os.path.normcase(os.path.normpath(root))
        return p == r or p.startswith(r.rstrip(os.sep) + os.sep)

    def _validate_path(self, path: str) -> str:
        """Validate and normalize path. Raises ValueError if blocked."""
        # os.path.expandvars is deliberately NOT called here. "%NAME%" is a legal
        # directory name on Windows and "$NAME"/"${NAME}" are legal on Linux and
        # macOS, so expanding them retargeted the operation at a DIFFERENT,
        # existing object without telling anyone: delete() on a folder literally
        # named "%USERNAME%" returned {"status": "deleted", ...\ADMIN} — the
        # sibling data folder was destroyed, the selected one survived, and the
        # API reported success (reproduced). Every caller — the WebUI, the chat
        # /reveal endpoint, cleanup and disk usage — hands over a real filesystem
        # path, never a shell string, so there was never anything to expand.
        expanded = os.path.expanduser(path)
        normalized = os.path.normpath(os.path.abspath(expanded))

        # "~/..." survives because SKILL.md tells the model to address the sandbox
        # that way. "~name" does not: it is indistinguishable from a real
        # directory of that name, and expanduser rewrites it to someone's home
        # directory. Guessing wrong there is the same silent-retarget failure, so
        # name both paths and refuse instead of picking one.
        if expanded != path and not (path == "~" or path[:2] in ("~/", "~" + os.sep)):
            raise ValueError(
                f"Từ chối vì đường dẫn bị diễn giải thành một đường dẫn khác:\n"
                f"  bạn yêu cầu: {path}\n"
                f"  hệ thống hiểu thành: {expanded}\n"
                f"Hãy dùng đường dẫn tuyệt đối đầy đủ."
            )
        # Resolve symlinks/junctions too: a link inside an allowed root must
        # not be a way to reach a blocked one.
        try:
            resolved = os.path.realpath(normalized)
        except OSError:
            resolved = normalized

        # Check blocked paths
        for blocked in BLOCKED_PATHS:
            if self._under(normalized, blocked) or self._under(resolved, blocked):
                raise ValueError(f"Đường dẫn bị chặn vì lý do bảo mật: {path}")

        # The data dir is inside the AI's allowlist, but not all of it: the
        # group manifests live there (AI_PROTECTED_DATA_SUBDIRS). Checked
        # before the allowlist so the refusal reads as a security block, not
        # as "outside the roots" with the roots listed.
        if self.enforce_roots:
            for blocked in self.ai_blocked:
                if self._under(normalized, blocked) or self._under(resolved, blocked):
                    raise ValueError(f"Đường dẫn bị chặn vì lý do bảo mật: {path}")

        # Check if within allowed roots — skipped for the human-facing UI
        # service (enforce_roots=False); the blocklist above already ran.
        if self.enforce_roots:
            in_allowed = any(
                self._under(normalized, root) and self._under(resolved, root)
                for root in self.allowed_roots
            )
            if not in_allowed:
                raise ValueError(
                    f"Đường dẫn nằm ngoài vùng cho phép: {path}\n"
                    f"Vùng cho phép: {', '.join(self.allowed_roots)}"
                )

        return normalized

    def validate_path(self, path: str) -> str:
        """Public alias for _validate_path. cleanup.py and disk_usage.py probe
        for a public `validate_path` first (getattr chain), so exposing it here
        lets sibling modules stop reaching into a private method."""
        return self._validate_path(path)

    def _file_info(self, path: str) -> Dict[str, Any]:
        """Get file/folder info dict.

        `path` keeps its exact OS spelling for os.stat here; only the two strings
        that end up in the JSON response are passed through _safe_str, because
        every caller of this method (list_dir, search, info, GET /info) feeds the
        result straight to the response encoder.
        """
        stat = os.stat(path)
        return {
            "name": _safe_str(os.path.basename(path)),
            "path": _safe_str(path),
            "is_dir": os.path.isdir(path),
            "size": stat.st_size if os.path.isfile(path) else 0,
            "size_human": self._human_size(stat.st_size) if os.path.isfile(path) else "",
            "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_mtime)),
            "created": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(stat.st_ctime)),
            "extension": os.path.splitext(path)[1].lower() if os.path.isfile(path) else "",
        }

    @staticmethod
    def _human_size(size_bytes: int) -> str:
        """Convert bytes to human readable string."""
        for unit in ("B", "KB", "MB", "GB"):
            if size_bytes < 1024:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024
        return f"{size_bytes:.1f} TB"

    # ── Core Operations ──────────────────────────────────────

    def list_dir(self, path: str, show_hidden: bool = False) -> Dict[str, Any]:
        """List contents of a directory."""
        safe_path = self._validate_path(path)
        if not os.path.isdir(safe_path):
            raise FileNotFoundError(f"Thư mục không tồn tại: {path}")

        items = []
        for entry in sorted(os.listdir(safe_path)):
            if not show_hidden and entry.startswith("."):
                continue
            full = os.path.join(safe_path, entry)
            try:
                items.append(self._file_info(full))
            except (PermissionError, OSError):
                # Same repair as _file_info: this branch is the one an
                # undecodable name is most likely to reach, so leaving it raw
                # would still poison the whole listing at encode time.
                items.append({"name": _safe_str(entry), "path": _safe_str(full),
                              "error": "Permission denied"})

        return {
            "path": safe_path,
            "parent": os.path.dirname(safe_path),
            "items": items,
            "count": len(items),
            "dirs": sum(1 for i in items if i.get("is_dir")),
            "files": sum(1 for i in items if not i.get("is_dir") and "error" not in i),
        }

    def create_folder(self, path: str) -> Dict[str, Any]:
        """Create a folder (including parents)."""
        safe_path = self._validate_path(path)
        if os.path.exists(safe_path):
            return {"status": "exists", "path": safe_path, "message": f"Thư mục đã tồn tại: {path}"}
        os.makedirs(safe_path, exist_ok=True)
        return {"status": "created", "path": safe_path, "message": f"Đã tạo thư mục: {safe_path}"}

    def create_file(self, path: str, content: str = "") -> Dict[str, Any]:
        """Create a file with optional content. Supports txt, docx, xlsx."""
        safe_path = self._validate_path(path)
        parent_dir = os.path.dirname(safe_path)
        os.makedirs(parent_dir, exist_ok=True)

        ext = os.path.splitext(safe_path)[1].lower()
        
        try:
            if ext == '.docx':
                from docx import Document
                doc = Document()
                if content:
                    doc.add_paragraph(content)
                doc.save(safe_path)
                size = os.path.getsize(safe_path)
            elif ext == '.xlsx':
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                if content:
                    for row_idx, line in enumerate(content.split('\n'), 1):
                        cells = line.split('\t') if '\t' in line else line.split(',')
                        for col_idx, cell in enumerate(cells, 1):
                            ws.cell(row=row_idx, column=col_idx, value=cell.strip())
                wb.save(safe_path)
                size = os.path.getsize(safe_path)
            else:
                mode = "w" if not os.path.exists(safe_path) else "w"
                with open(safe_path, mode, encoding="utf-8") as f:
                    f.write(content)
                size = os.path.getsize(safe_path)
                
            return {
                "status": "created",
                "path": safe_path,
                "size": size,
                "extension": ext,
                "message": f"Đã tạo file: {safe_path}",
            }
        except ImportError as e:
            raise RuntimeError(f"Thiếu thư viện để xử lý file {ext}: {str(e)}")
        except Exception as e:
            raise RuntimeError(f"Lỗi tạo file {ext}: {str(e)}")

    def read_file(self, path: str, max_lines: int = 1000) -> Dict[str, Any]:
        """Read file content (supports txt, docx, xlsx) and metadata."""
        safe_path = self._validate_path(path)
        if not os.path.isfile(safe_path):
            raise FileNotFoundError(f"File không tồn tại: {path}")

        size = os.path.getsize(safe_path)
        if size > MAX_FILE_SIZE_MB * 1024 * 1024:
            return {"error": f"File quá lớn ({self._human_size(size)}). Giới hạn: {MAX_FILE_SIZE_MB}MB"}

        ext = os.path.splitext(safe_path)[1].lower()
        content = ""
        lines = []

        try:
            if ext == '.docx':
                from docx import Document
                doc = Document(safe_path)
                for i, p in enumerate(doc.paragraphs):
                    if i >= max_lines: break
                    if p.text.strip():
                        lines.append(p.text)
                content = "\n".join(lines)
            elif ext == '.xlsx':
                from openpyxl import load_workbook
                wb = None
                try:
                    wb = load_workbook(safe_path, read_only=True, data_only=True)
                    for ws in wb.worksheets:
                        lines.append(f"--- Sheet: {ws.title} ---")
                        for i, row in enumerate(ws.iter_rows(values_only=True)):
                            if len(lines) >= max_lines: break
                            row_data = [str(c) if c is not None else "" for c in row]
                            if any(row_data):
                                lines.append("\t".join(row_data))
                        if len(lines) >= max_lines: break
                    content = "\n".join(lines)
                finally:
                    if wb:
                        wb.close()
            else:
                with open(safe_path, "r", encoding="utf-8") as f:
                    for i, line in enumerate(f):
                        if i >= max_lines:
                            break
                        lines.append(line.rstrip("\n"))
                content = "\n".join(lines)
        except UnicodeDecodeError:
            return {
                "path": safe_path,
                "is_binary": True,
                "size": size,
                "size_human": self._human_size(size),
                "extension": ext,
                "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(safe_path))),
                "message": "File nhị phân, không thể đọc dạng text.",
            }
        except ImportError as e:
            return {"error": f"Thiếu thư viện để đọc file {ext}: {str(e)}"}
        except Exception as e:
            return {"error": f"Lỗi đọc file: {str(e)}"}

        return {
            "path": safe_path,
            "content": content,
            "lines": len(lines),
            "truncated": len(lines) >= max_lines,
            "size": size,
            "size_human": self._human_size(size),
            "extension": ext,
            "modified": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getmtime(safe_path))),
            "created": time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(os.path.getctime(safe_path))),
        }

    def delete(self, path: str) -> Dict[str, Any]:
        """Delete a file or folder."""
        safe_path = self._validate_path(path)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Đường dẫn không tồn tại: {path}")

        if os.path.isdir(safe_path):
            item_count = sum(len(files) + len(dirs) for _, dirs, files in os.walk(safe_path))
            shutil.rmtree(safe_path)
            return {"status": "deleted", "path": safe_path, "type": "folder", "items_removed": item_count}
        else:
            os.remove(safe_path)
            return {"status": "deleted", "path": safe_path, "type": "file"}

    def move(self, src: str, dst: str) -> Dict[str, Any]:
        """Move or rename a file/folder."""
        safe_src = self._validate_path(src)
        safe_dst = self._validate_path(dst)
        if not os.path.exists(safe_src):
            raise FileNotFoundError(f"Nguồn không tồn tại: {src}")

        # If dst is a directory, move into it
        if os.path.isdir(safe_dst):
            safe_dst = os.path.join(safe_dst, os.path.basename(safe_src))

        os.makedirs(os.path.dirname(safe_dst), exist_ok=True)
        shutil.move(safe_src, safe_dst)
        return {"status": "moved", "from": safe_src, "to": safe_dst}

    def copy(self, src: str, dst: str) -> Dict[str, Any]:
        """Copy a file or folder."""
        safe_src = self._validate_path(src)
        safe_dst = self._validate_path(dst)
        if not os.path.exists(safe_src):
            raise FileNotFoundError(f"Nguồn không tồn tại: {src}")

        os.makedirs(os.path.dirname(safe_dst), exist_ok=True)

        if os.path.isdir(safe_src):
            shutil.copytree(safe_src, safe_dst)
        else:
            shutil.copy2(safe_src, safe_dst)
        return {"status": "copied", "from": safe_src, "to": safe_dst}

    def info(self, path: str) -> Dict[str, Any]:
        """Get detailed file/folder info."""
        safe_path = self._validate_path(path)
        if not os.path.exists(safe_path):
            raise FileNotFoundError(f"Đường dẫn không tồn tại: {path}")

        info = self._file_info(safe_path)
        if info["is_dir"]:
            # Count contents
            total_files = 0
            total_size = 0
            for root, dirs, files in os.walk(safe_path):
                total_files += len(files)
                for f in files:
                    try:
                        total_size += os.path.getsize(os.path.join(root, f))
                    except OSError:
                        pass
            info["total_files"] = total_files
            info["total_size"] = total_size
            info["total_size_human"] = self._human_size(total_size)
        return info

    def search(self, path: str, pattern: str = "*", recursive: bool = True) -> Dict[str, Any]:
        """Search files matching pattern (supports mutiple patterns separated by |)."""
        safe_path = self._validate_path(path)
        if not os.path.isdir(safe_path):
            raise FileNotFoundError(f"Thư mục không tồn tại: {path}")

        matches = []
        patterns = pattern.split("|") if "|" in pattern else [pattern]
        
        seen_paths = set()
        for p in patterns:
            search_pattern = os.path.join(safe_path, "**", p.strip()) if recursive else os.path.join(safe_path, p.strip())
            for match in glob.iglob(search_pattern, recursive=recursive):
                if match in seen_paths: continue
                seen_paths.add(match)
                try:
                    matches.append(self._file_info(match))
                except (PermissionError, OSError):
                    continue
                if len(matches) >= 200:
                    break
            if len(matches) >= 200:
                break

        return {"path": safe_path, "pattern": pattern, "matches": matches, "count": len(matches)}

    def write_bytes(self, path: str, data: bytes) -> Dict[str, Any]:
        """Write raw bytes to a file (Drive downloads land through here).

        create_file() decodes/encodes text and knows .docx/.xlsx; it is wrong
        for arbitrary binary payloads, so this is a separate primitive rather
        than a mode flag on it.
        """
        safe_path = self._validate_path(path)
        os.makedirs(os.path.dirname(safe_path), exist_ok=True)
        with open(safe_path, "wb") as f:
            f.write(data)
        size = os.path.getsize(safe_path)
        return {
            "status": "created",
            "path": safe_path,
            "size": size,
            "size_human": self._human_size(size),
        }

    # ── Trình sửa có cấu trúc (canvas): xlsx dạng lưới, docx dạng đoạn có format ──
    # read_file() trả TEXT phẳng (cho AI/xem nhanh); các hàm dưới giữ CẤU TRÚC để sửa rồi
    # ghi lại ĐÚNG định dạng gốc, thay vì làm phẳng thành .txt.

    MAX_SHEET_ROWS = 500
    MAX_SHEET_COLS = 60

    def read_sheet(self, path: str) -> Dict[str, Any]:
        """Đọc .xlsx thành các sheet dạng lưới ô (giới hạn 500 dòng × 60 cột mỗi sheet)."""
        safe_path = self._validate_path(path)
        if not os.path.isfile(safe_path):
            raise FileNotFoundError(f"File không tồn tại: {path}")
        from openpyxl import load_workbook
        wb = load_workbook(safe_path, read_only=True, data_only=True)
        try:
            sheets = []
            for ws in wb.worksheets:
                rows, truncated = [], False
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= self.MAX_SHEET_ROWS:
                        truncated = True
                        break
                    cells = ["" if c is None else str(c) for c in row[: self.MAX_SHEET_COLS]]
                    rows.append(cells)
                width = max([len(r) for r in rows] or [1])
                for r in rows:
                    r.extend([""] * (width - len(r)))
                sheets.append({"name": ws.title, "rows": rows, "truncated": truncated})
            return {"sheets": sheets, "path": safe_path}
        finally:
            wb.close()

    def write_sheet(self, path: str, sheets: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Ghi lại .xlsx từ các sheet dạng lưới. GHI ĐÈ nội dung ô, giữ đúng định dạng file."""
        safe_path = self._validate_path(path)
        from openpyxl import Workbook
        wb = Workbook()
        wb.remove(wb.active)
        for sh in (sheets or []):
            ws = wb.create_sheet(title=(str(sh.get("name") or "Sheet")[:31] or "Sheet"))
            for row in (sh.get("rows") or []):
                ws.append(["" if c is None else c for c in row])
        if not wb.sheetnames:
            wb.create_sheet(title="Sheet")
        wb.save(safe_path)
        size = os.path.getsize(safe_path)
        return {"status": "saved", "path": safe_path, "size": size, "size_human": self._human_size(size)}

    # ── Sửa bảng tính TẠI CHỖ (xlsx_read / xlsx_append / xlsx_write) ──────────
    # write_sheet() ở trên dựng lại cả workbook từ lưới ô: tab không gửi lên biến
    # mất, định dạng/độ rộng cột/công thức cũng vậy. Agent trong nhóm chỉ được
    # nối dòng hoặc sửa vài ô trong file MẪU của chủ (template upload video…),
    # nên các hàm dưới load_workbook (data_only=False để giữ công thức) rồi chỉ
    # chạm đúng ô cần đổi. Giới hạn của openpyxl vẫn còn: ảnh/biểu đồ nhúng trong
    # workbook không sống sót qua load→save — nói rõ với người dùng khi họ hỏi.
    # CSV không có tab/ô nên chỉ đọc + nối dòng; muốn sửa ô thì chuyển sang .xlsx.

    SPREADSHEET_EXTS = (".xlsx", ".xlsm", ".csv", ".tsv")
    _CELL_REF = re.compile(r"^([A-Za-z]{1,3})([1-9][0-9]{0,6})$")

    @staticmethod
    def _cell_value(v: Any) -> Any:
        """Giá trị từ JSON của model → giá trị ô. openpyxl chỉ nhận scalar; list/dict
        lồng nhau thì ghi dạng chuỗi còn hơn nổ giữa chừng sau khi đã ghi nửa bảng."""
        if v is None or isinstance(v, (str, int, float, bool)):
            return v
        return str(v)

    @classmethod
    def _norm_rows(cls, rows: Any) -> List[List[Any]]:
        """Chuẩn hoá 'rows' model gửi: [[a,b],[c,d]] là chuẩn; một list phẳng [a,b]
        được hiểu là MỘT dòng (model hay gửi thế khi chỉ thêm một bản ghi)."""
        if rows is None:
            raise ValueError("Thiếu 'rows' (danh sách dòng, mỗi dòng là một danh sách ô)")
        if isinstance(rows, (str, bytes, dict)) or not hasattr(rows, "__iter__"):
            raise ValueError("'rows' phải là danh sách các dòng, ví dụ [[\"a\", 1], [\"b\", 2]]")
        rows = list(rows)
        if rows and not any(isinstance(r, (list, tuple, dict)) for r in rows):
            return [[cls._cell_value(c) for c in rows]]
        out = []
        for r in rows:
            if isinstance(r, dict):
                r = list(r.values())
            elif not isinstance(r, (list, tuple)):
                r = [r]
            out.append([cls._cell_value(c) for c in r])
        return out

    def _spreadsheet_path(self, path: str, exts=SPREADSHEET_EXTS):
        """Sandbox + tồn tại + đúng loại file. Trả (safe_path, ext)."""
        safe_path = self._validate_path(path)
        if not os.path.isfile(safe_path):
            raise FileNotFoundError(f"File không tồn tại: {path}")
        ext = os.path.splitext(safe_path)[1].lower()
        if ext not in exts:
            raise ValueError(f"Không phải file bảng tính ({', '.join(exts)}): {path}")
        return safe_path, ext

    @staticmethod
    def _csv_delim(ext: str) -> str:
        return "\t" if ext == ".tsv" else ","

    @staticmethod
    def _pick_ws(wb, sheet: Any):
        """Chọn tab: tên (khớp chính xác, rồi bỏ hoa/thường + khoảng trắng), hoặc số
        thứ tự 1-based; rỗng = tab đang mở. Không thấy → ValueError kể tên các tab
        để model tự sửa thay vì đoán."""
        names = wb.sheetnames
        if sheet is None or str(sheet).strip() == "":
            return wb.active if wb.active is not None else wb.worksheets[0]
        key = str(sheet).strip()
        if key in names:
            return wb[key]
        for n in names:
            if n.casefold().strip() == key.casefold():
                return wb[n]
        if key.isdigit() and 1 <= int(key) <= len(names):
            return wb.worksheets[int(key) - 1]
        raise ValueError(f"Không có sheet '{sheet}'. Các sheet hiện có: {', '.join(names)}")

    @staticmethod
    def _last_used_row(ws) -> int:
        """Dòng cuối CÓ dữ liệu. ws.max_row tính cả dòng trống đã kẻ viền sẵn trong
        file mẫu; dựa vào nó thì dòng mới bị đẩy xuống dưới một khoảng trắng."""
        last = 0
        for idx, row in enumerate(ws.iter_rows(values_only=True), 1):
            if any(c is not None and str(c) != "" for c in row):
                last = idx
        return last

    @staticmethod
    def _strip_trailing_blank(rows: List[List[Any]]) -> List[List[Any]]:
        while rows and not any(c not in (None, "") for c in rows[-1]):
            rows.pop()
        return rows

    @staticmethod
    def _formula_text(v: Any) -> Optional[str]:
        """Công thức của ô khi mở data_only=False ('=SUM(A1:A2)'), None nếu là ô
        thường. Công thức mảng / bảng dữ liệu về dạng object có .text."""
        if isinstance(v, str):
            return v if v.startswith("=") else None
        t = getattr(v, "text", None)
        if isinstance(t, str) and t:
            return t if t.startswith("=") else "=" + t
        return None

    @staticmethod
    def _rtrim_cells(row: list) -> list:
        """Bo cac o rong (None / chuoi trang) o cuoi dong; giu o rong nam giua."""
        end = len(row)
        while end and (row[end - 1] is None or (isinstance(row[end - 1], str) and not row[end - 1].strip())):
            end -= 1
        return row[:end]

    def _xlsx_rows(self, safe_path: str, sheet: Any, max_cols: int):
        """(tên tab, danh sách tab, dòng, số ô công thức, [(dòng, cột) ô công thức
        không có kết quả]).

        openpyxl không tính công thức, và khi append_sheet_rows / update_sheet_cells
        lưu lại thì kết quả Excel đã cache cũng không được ghi ra — nên mở
        data_only=True sau lần sửa đầu tiên trả None cho MỌI ô công thức: tổng,
        lookup trong file mẫu của chủ hiện thành ô trống và model tin đó là
        trống thật. Vì vậy đọc công thức trước; tab nào có công thức thì mở thêm
        một lần data_only=True: ô còn cache lấy giá trị, ô mất cache hiện chính
        công thức và được đếm để handler nói rõ con số chưa được tính."""
        from openpyxl import load_workbook
        wb = load_workbook(safe_path, read_only=True, data_only=False)
        try:
            ws = self._pick_ws(wb, sheet)
            title, names = ws.title, list(wb.sheetnames)
            rows = [list(r) for r in ws.iter_rows(values_only=True, max_col=max_cols)]
        finally:
            wb.close()
        # max_col lam openpyxl dem du max_cols o cho MOI dong, ke ca o chua tung co gia
        # tri: bang tra ve model thanh "| a | b |  |  | ... |" voi hang chuc o rong moi
        # dong — toan ngu canh ma khong mang thong tin. Cat o rong o duoi cung moi dong.
        rows = [self._rtrim_cells(r) for r in rows]
        formulas = []
        for i, r in enumerate(rows):
            for j, c in enumerate(r):
                f = self._formula_text(c)
                if f is not None:
                    formulas.append((i, j, f))
        uncomputed: List[tuple] = []
        if formulas:
            wb = load_workbook(safe_path, read_only=True, data_only=True)
            try:
                cached = [list(r) for r in wb[title].iter_rows(values_only=True, max_col=max_cols)]
            finally:
                wb.close()
            for i, j, f in formulas:
                v = cached[i][j] if i < len(cached) and j < len(cached[i]) else None
                if v is None:
                    uncomputed.append((i, j))
                rows[i][j] = f if v is None else v
        return title, names, rows, len(formulas), uncomputed

    def _csv_rows(self, safe_path: str, ext: str, max_cols: int) -> List[List[str]]:
        # utf-8-sig: Excel ghi BOM ở đầu; errors=replace: một ô lệch mã không được
        # làm hỏng cả lần đọc (cùng lý do _safe_str ở trên).
        with open(safe_path, "r", newline="", encoding="utf-8-sig", errors="replace") as f:
            return [self._rtrim_cells(list(r[:max_cols])) for r in csv.reader(f, delimiter=self._csv_delim(ext))]

    def read_sheet_rows(self, path: str, sheet: Any = None, max_rows: int = 100,
                        max_cols: Optional[int] = None) -> Dict[str, Any]:
        """Đọc MỘT tab (hoặc CSV) thành danh sách dòng chuỗi, cắt ở max_rows nhưng vẫn
        báo tổng số dòng để model biết còn bao nhiêu chưa xem."""
        safe_path, ext = self._spreadsheet_path(path)
        try:
            max_rows = int(max_rows or 100)
        except (TypeError, ValueError):
            max_rows = 100
        max_rows = max(1, min(max_rows, self.MAX_SHEET_ROWS))
        max_cols = max(1, min(int(max_cols or self.MAX_SHEET_COLS), self.MAX_SHEET_COLS))

        if ext in (".csv", ".tsv"):
            rows = self._strip_trailing_blank(self._csv_rows(safe_path, ext, max_cols))
            name, names, n_formula, uncomputed = "", [], 0, []
        else:
            name, names, rows, n_formula, uncomputed = self._xlsx_rows(safe_path, sheet, max_cols)
            rows = self._strip_trailing_blank(rows)

        total = len(rows)
        shown = [["" if c is None else str(c) for c in r] for r in rows[:max_rows]]
        width = max([len(r) for r in shown] or [1])
        for r in shown:
            r.extend([""] * (width - len(r)))
        return {
            "path": safe_path,
            "sheet": name,
            "sheets": names,
            "rows": shown,
            "total_rows": total,
            "max_rows": max_rows,
            "truncated": total > len(shown),
            "formula_cells": n_formula,
            # Chỉ đếm trong phần đang hiện: ghi chú đi kèm bảng nói về bảng đó.
            "formulas_uncomputed": sum(1 for i, _ in uncomputed if i < max_rows),
        }

    def append_sheet_rows(self, path: str, sheet: Any = None, rows: Any = None) -> Dict[str, Any]:
        """Nối dòng vào SAU dòng cuối có dữ liệu của một tab (hoặc cuối file CSV),
        giữ nguyên mọi tab khác, định dạng và công thức đang có."""
        safe_path, ext = self._spreadsheet_path(path)
        rows = self._norm_rows(rows)
        if not rows:
            raise ValueError("'rows' rỗng — không có gì để thêm")

        if ext in (".csv", ".tsv"):
            delim = self._csv_delim(ext)
            existing = self._strip_trailing_blank(self._csv_rows(safe_path, ext, self.MAX_SHEET_COLS))
            first = len(existing) + 1
            # File do tay người/ứng dụng khác ghi có thể thiếu newline cuối; nối thẳng
            # vào thì dòng mới dính vào dòng cũ thành một bản ghi hỏng.
            with open(safe_path, "rb+") as f:
                f.seek(0, os.SEEK_END)
                if f.tell() > 0:
                    f.seek(-1, os.SEEK_END)
                    if f.read(1) not in (b"\n", b"\r"):
                        f.write(b"\r\n")
            with open(safe_path, "a", newline="", encoding="utf-8") as f:
                w = csv.writer(f, delimiter=delim)
                for r in rows:
                    w.writerow(["" if c is None else c for c in r])
            return {"status": "appended", "path": safe_path, "sheet": "", "rows_added": len(rows),
                    "first_row": first, "last_row": first + len(rows) - 1}

        from openpyxl import load_workbook
        wb = load_workbook(safe_path, keep_vba=(ext == ".xlsm"))
        try:
            ws = self._pick_ws(wb, sheet)
            first = self._last_used_row(ws) + 1
            for i, r in enumerate(rows):
                for j, v in enumerate(r):
                    ws.cell(row=first + i, column=j + 1, value=v)
            wb.save(safe_path)
            title = ws.title
        finally:
            wb.close()
        return {"status": "appended", "path": safe_path, "sheet": title, "rows_added": len(rows),
                "first_row": first, "last_row": first + len(rows) - 1}

    def update_sheet_cells(self, path: str, sheet: Any = None, cells: Optional[Dict[str, Any]] = None,
                           rows: Any = None, start: str = "A1") -> Dict[str, Any]:
        """Ghi đè từng ô ({"A1": v, "B2": 3}) và/hoặc một khối 'rows' bắt đầu từ ô
        'start'; mọi ô khác, tab khác, định dạng, công thức giữ nguyên."""
        safe_path, ext = self._spreadsheet_path(path, exts=(".xlsx", ".xlsm"))
        if not cells and not rows:
            raise ValueError("Cần 'cells' ({\"A1\": ...}) hoặc 'rows' + 'start' để biết ghi gì vào đâu")
        if cells is not None and not isinstance(cells, dict):
            raise ValueError("'cells' phải là object {\"A1\": giá_trị, ...}")

        from openpyxl import load_workbook
        from openpyxl.utils.cell import column_index_from_string, get_column_letter
        wb = load_workbook(safe_path, keep_vba=(ext == ".xlsm"))
        try:
            ws = self._pick_ws(wb, sheet)
            written: List[str] = []
            block_range = ""
            for ref, val in (cells or {}).items():
                m = self._CELL_REF.match(str(ref).strip())
                if not m:
                    raise ValueError(f"Ô không hợp lệ: '{ref}' (dạng A1, B12, AC7)")
                coord = m.group(1).upper() + m.group(2)
                ws[coord] = self._cell_value(val)
                written.append(coord)
            if rows:
                block = self._norm_rows(rows)
                m = self._CELL_REF.match(str(start or "A1").strip())
                if not m:
                    raise ValueError(f"'start' không hợp lệ: '{start}' (dạng A1)")
                col0, row0 = column_index_from_string(m.group(1).upper()), int(m.group(2))
                width = 0
                for i, r in enumerate(block):
                    width = max(width, len(r))
                    for j, v in enumerate(r):
                        ws.cell(row=row0 + i, column=col0 + j, value=v)
                if block and width:
                    block_range = (f"{m.group(1).upper()}{row0}:"
                                   f"{get_column_letter(col0 + width - 1)}{row0 + len(block) - 1}")
                    written.append(block_range)
            wb.save(safe_path)
            title = ws.title
        finally:
            wb.close()
        return {"status": "saved", "path": safe_path, "sheet": title,
                "cells_written": len(cells or {}), "range": block_range, "targets": written[:50]}

    # Căn lề: python-docx dùng enum WD_PARAGRAPH_ALIGNMENT (0..3). Văn bản hành chính VN
    # phụ thuộc nặng vào CĂN GIỮA (quốc hiệu, tiêu đề) nên phải đọc/ghi được, nếu không
    # tài liệu mở ra trông "mất format" dù chữ vẫn đúng.
    _ALIGN_TO_STR = {0: "left", 1: "center", 2: "right", 3: "justify"}
    _STR_TO_ALIGN = {"left": 0, "center": 1, "right": 2, "justify": 3}

    def read_doc(self, path: str) -> Dict[str, Any]:
        """Đọc .docx thành các đoạn kèm style, căn lề, cỡ chữ, đậm/nghiêng/gạch chân."""
        safe_path = self._validate_path(path)
        if not os.path.isfile(safe_path):
            raise FileNotFoundError(f"File không tồn tại: {path}")
        from docx import Document
        doc = Document(safe_path)

        def _base_pt(p):
            """Cỡ chữ của đoạn: ưu tiên run có chữ, rồi tới style, rồi mặc định của tài liệu."""
            for r in p.runs:
                if r.text.strip() and r.font.size is not None:
                    return round(r.font.size.pt, 1)
            try:
                if p.style is not None and p.style.font.size is not None:
                    return round(p.style.font.size.pt, 1)
            except Exception:
                pass
            try:
                sz = doc.styles["Normal"].font.size
                if sz is not None:
                    return round(sz.pt, 1)
            except Exception:
                pass
            return None

        paras = []
        for p in doc.paragraphs:
            style = (p.style.name if p.style is not None else "") or "Normal"
            marked = [r for r in p.runs if r.text.strip()]
            align = p.alignment
            if align is None:                       # kế thừa từ style của đoạn
                try:
                    align = p.style.paragraph_format.alignment
                except Exception:
                    align = None
            paras.append({
                "text": p.text,
                "style": style,
                # Đậm/nghiêng/gạch chân ở mức ĐOẠN (đủ cho trình sửa nhẹ).
                "bold": bool(marked and all(r.bold for r in marked)),
                "italic": bool(marked and all(r.italic for r in marked)),
                "underline": bool(marked and all(r.underline for r in marked)),
                "align": self._ALIGN_TO_STR.get(int(align), "left") if align is not None else None,
                "size": _base_pt(p),
            })
        return {"paragraphs": paras, "path": safe_path}

    def write_doc(self, path: str, paragraphs: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Ghi lại .docx từ danh sách đoạn (style, căn lề, cỡ chữ, đậm/nghiêng/gạch chân)."""
        safe_path = self._validate_path(path)
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        for p in (paragraphs or []):
            style = str(p.get("style") or "Normal")
            text = "" if p.get("text") is None else str(p.get("text"))
            try:
                para = doc.add_paragraph(style=style)
            except Exception:
                para = doc.add_paragraph()   # style lạ (file gốc dùng style riêng) → dùng mặc định
            al = p.get("align")
            if al in self._STR_TO_ALIGN:
                para.alignment = self._STR_TO_ALIGN[al]
            run = para.add_run(text)
            if p.get("bold"):
                run.bold = True
            if p.get("italic"):
                run.italic = True
            if p.get("underline"):
                run.underline = True
            try:
                if p.get("size"):
                    run.font.size = Pt(float(p["size"]))
            except Exception:
                pass
        doc.save(safe_path)
        size = os.path.getsize(safe_path)
        return {"status": "saved", "path": safe_path, "size": size, "size_human": self._human_size(size)}

    def get_allowed_roots(self) -> List[Dict[str, str]]:
        """Return list of quick-access root directories for the UI.

        With enforce_roots=False these are shortcuts, not a boundary, so the
        home directory is prepended — it is the natural place to start browsing
        once browsing is no longer fenced in.
        """
        listing = list(self.allowed_roots)
        if not self.enforce_roots:
            home = os.path.normpath(os.path.expanduser("~"))
            if home not in listing:
                listing.insert(0, home)
        roots = []
        for r in listing:
            if os.path.isdir(r):
                roots.append({"path": r, "name": os.path.basename(r) or r, "exists": True})
            else:
                roots.append({"path": r, "name": os.path.basename(r) or r, "exists": False})
        return roots


# Global singleton — sandboxed. Every AI-facing caller (brain file_action, the
# file_manager workflow node, Telegram actions, chat pipeline) imports this one.
file_service = FileService()

# The File Manager web UI's service: same blocklist, no allowlist. Only
# routes.py should import this — a logged-in human browsing their own machine
# is not the threat the sandbox exists for.
user_file_service = FileService(enforce_roots=False)
